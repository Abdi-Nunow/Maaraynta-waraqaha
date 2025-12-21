import streamlit as st
import pandas as pd
import io
from datetime import datetime
from docx import Document
import base64
import os

# ================= PAGE CONFIG =================
st.set_page_config(page_title="Maaraynta Waraaqaha", layout="wide")

# ================= LOGO =================
if os.path.exists("images.png"):
    st.image("images.png", width=200)

st.markdown("## 📁 Nidaamka Maareynta Waraaqaha")
st.markdown("Waxaa loogu talagalay in waaxyaha kala duwan ee xafiiska dakhliga ay isku diraan waraaqaha.")

# ================= FILES =================
passwords_file = "passwords.csv"
waraaqaha_file = "waraaqaha.csv"

# ================= DEFAULT WAAXYO + PASSWORDS =================
if not os.path.exists(passwords_file):
    default_passwords = {
        "Xafiiska Wasiirka": "Admin2100",
        "Wasiir Ku-xigeenka 1aad": "Admin2100",
        "Wasiir Ku-xigeenka 2aad": "Admin2100",
        "Wasiir Ku-xigeenka 3aad": "Admin2100",
        "Secretory": "Admin2100",
        "Waaxda Auditka": "Admin2100",
        "Waaxda ICT": "Admin2100",
        "Waaxda HRM": "Admin2100",
        "Waaxda Public Relation": "Admin2100",
        "Waaxda Adeega Shacabka": "Admin2100",
        "Arkiviya-1": "Admin2100",
        "Arkiviya-2": "Admin2100"
    }
    pd.DataFrame(default_passwords.items(), columns=["waaxda", "password"]).to_csv(passwords_file, index=False)

df_passwords = pd.read_csv(passwords_file)
waaxyo_passwords = dict(zip(df_passwords.waaxda, df_passwords.password))

# ================= ADMIN =================
ADMIN_USER = "Admin"
ADMIN_PASS = "Admin2100"

# ================= SESSION =================
if "user" not in st.session_state:
    st.session_state.user = None
    st.session_state.is_admin = False

# ================= LOGIN =================
if st.session_state.user is None:
    st.subheader("🔐 Login")
    login_type = st.radio("Nooca isticmaalaha", ["Waax", "Admin"])

    if login_type == "Waax":
        waax = st.selectbox("Dooro Waaxda", list(waaxyo_passwords.keys()))
        pwd = st.text_input("Password", type="password")
        if st.button("Gali"):
            if pwd == waaxyo_passwords.get(waax):
                st.session_state.user = waax
                st.experimental_rerun()
            else:
                st.error("Password khaldan ❌")

    else:
        u = st.text_input("Admin Username")
        p = st.text_input("Admin Password", type="password")
        if st.button("Gali"):
            if u == ADMIN_USER and p == ADMIN_PASS:
                st.session_state.user = "Admin"
                st.session_state.is_admin = True
                st.experimental_rerun()
            else:
                st.error("Admin login khaldan ❌")

# ================= MAIN APP =================
else:
    user = st.session_state.user
    is_admin = st.session_state.is_admin
    st.success(f"Ku soo dhawoow: {user}")

    if os.path.exists(waraaqaha_file):
        df = pd.read_csv(waraaqaha_file)
    else:
        df = pd.DataFrame(columns=[
            "Ka Socota", "Loogu Talagalay", "Cinwaan",
            "Qoraal", "Taariikh", "Diary", "File", "FileData"
        ])

    # ================= SEND LETTER =================
    st.subheader("📤 Dir Waraaq")
    col1, col2 = st.columns(2)

    with col1:
        cinwaan = st.text_input("Cinwaanka")
    with col2:
        loo = st.selectbox(
            "Loogu Talagalay",
            [w for w in waaxyo_passwords.keys() if w != user]
        )

    qoraal = st.text_area("Qoraalka Waraaqda")
    file = st.file_uploader("Lifaaq (ikhtiyaari)", type=["pdf", "docx", "xlsx"])

    fname, fdata = "", ""
    if file:
        fname = file.name
        fdata = base64.b64encode(file.read()).decode()

    if st.button("📨 Dir Waraaq"):
        diary_no = f"DR-{len(df)+1:05d}"
        new = {
            "Ka Socota": user,
            "Loogu Talagalay": loo,
            "Cinwaan": cinwaan,
            "Qoraal": qoraal,
            "Taariikh": datetime.today().strftime("%Y-%m-%d"),
            "Diary": diary_no,
            "File": fname,
            "FileData": fdata
        }
        df = pd.concat([df, pd.DataFrame([new])], ignore_index=True)
        df.to_csv(waraaqaha_file, index=False)
        st.success(f"✅ Waraaqda waa la diray | Diary: {diary_no}")

    # ================= VIEW LETTERS =================
    st.subheader("📥 Waraaqaha")
    view_df = df if is_admin else df[df["Loogu Talagalay"] == user]
    st.dataframe(view_df.drop(columns=["FileData"], errors="ignore"))

    # ================= DOWNLOADS =================
    if not view_df.empty:
        # WORD
        doc = Document()
        doc.add_heading("Waraaqaha", 0)
        for _, r in view_df.iterrows():
            doc.add_paragraph(f"Diary: {r['Diary']}")
            doc.add_paragraph(f"Taariikh: {r['Taariikh']}")
            doc.add_paragraph(f"Ka Socota: {r['Ka Socota']}")
            doc.add_paragraph(f"Cinwaan: {r['Cinwaan']}")
            doc.add_paragraph(r['Qoraal'])
            doc.add_paragraph("------")

        buf = io.BytesIO()
        doc.save(buf)

        st.download_button(
            "📄 Soo Degso Word",
            data=buf.getvalue(),
            file_name="waraaqaha.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="word_download"
        )

        # EXCEL
        excel_buf = io.BytesIO()
        view_df.drop(columns=["FileData"], errors="ignore").to_excel(excel_buf, index=False)

        st.download_button(
            "📊 Soo Degso Excel",
            data=excel_buf.getvalue(),
            file_name="waraaqaha.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="excel_download"
        )

    # ================= LOGOUT =================
    if st.button("🚪 Logout"):
        st.session_state.clear()
        st.experimental_rerun()
